"""Daily automation script (v2).

Cleans up notes, summarizes them via OpenAI, writes multiple output
formats, and optionally updates a GitHub issue for the day's run.
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

import requests
from docx import Document
from openai import OpenAI
from openpyxl import Workbook, load_workbook

BASE_DIR = Path(__file__).resolve().parent
NOTES_DIR = BASE_DIR / "notes"
OUTPUT_DIR = BASE_DIR / "output"
TXT_PATH = OUTPUT_DIR / "daily_summary.txt"
DOCX_PATH = OUTPUT_DIR / "daily_summary.docx"
XLSX_PATH = OUTPUT_DIR / "daily_log.xlsx"
JSON_PATH = OUTPUT_DIR / "daily_summary.json"
REPO_JSON_PATH = Path(os.getenv("NEXTJS_JSON_PATH") or (BASE_DIR / "data" / "daily_summary.json"))


def log(message: str) -> None:
    """Simple timestamped logger."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{now}] {message}", flush=True)


def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def check_openai_client() -> OpenAI:
    """Initialize and return an OpenAI client, or exit with a clear log."""
    log("Checking OPENAI_API_KEY...")
    try:
        client = OpenAI()
    except Exception as exc:  # noqa: BLE001
        log(f"Failed to initialize OpenAI client: {exc}")
        sys.exit(1)

    log("OpenAI client initialized.")
    return client


def load_recent_notes(max_files: int = 5) -> str:
    """
    Load the contents of the most recently modified note files (txt/md)
    from NOTES_DIR and concatenate them.
    """
    log("Loading recent notes...")

    if not NOTES_DIR.exists():
        log(f"No notes directory found at {NOTES_DIR}, using placeholder text.")
        return "No notes were found; this is a placeholder run."

    note_files = sorted(
        [p for p in NOTES_DIR.iterdir() if p.suffix in {".txt", ".md"}],
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )

    if not note_files:
        log("No .txt or .md files found in notes; using placeholder text.")
        return "No notes were found; this is a placeholder run."

    selected = note_files[:max_files]
    log(f"Using {len(selected)} note file(s):")
    for file_path in selected:
        log(f"  - {file_path.name}")

    chunks: list[str] = []
    for path in selected:
        try:
            chunks.append(f"# {path.name}\n")
            chunks.append(path.read_text(encoding="utf-8"))
            chunks.append("\n\n")
        except Exception as exc:  # noqa: BLE001
            log(f"WARNING: Could not read {path}: {exc}")

    combined = "\n".join(chunks).strip()
    if not combined:
        return "Notes existed but could not be read; this is a placeholder run."

    return combined


def summarize_notes(client: OpenAI, notes_text: str) -> str:
    """
    Call OpenAI Responses API to turn raw notes into a structured daily summary.
    Returns a plain text summary string.
    """
    log("Summarizing notes via OpenAI...")

    try:
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "user",
                    "content": (
                        "You are an assistant that turns raw scratchpad notes into a "
                        "structured daily summary for my automation/dev portfolio.\n\n"
                        "Return a single, dense text block — no markdown, no bullet "
                        "syntax. Include:\n"
                        "- One-paragraph high-level summary\n"
                        "- Key work completed\n"
                        "- Any blockers or TODOs\n\n"
                        f"Here are today's raw notes:\n\n{notes_text}"
                    ),
                }
            ],
        )
    except Exception as exc:  # noqa: BLE001
        log(f"OpenAI request failed: {exc}")
        raise

    chunks: list[str] = []
    for segment in getattr(response, "output", []):
        for part in getattr(segment, "content", []):
            if getattr(part, "type", None) == "output_text":
                chunks.append(part.text)

    summary = "\n".join(chunks).strip()
    if not summary:
        raise ValueError("OpenAI response did not contain summary text.")

    log("Got structured summary from OpenAI.")
    return summary


def write_txt_summary(summary: str) -> Path:
    log("Writing text summary...")
    ensure_output_dir()
    TXT_PATH.write_text(summary + "\n", encoding="utf-8")
    log(f"Text summary written to {TXT_PATH}")
    return TXT_PATH


def write_docx_summary(summary: str) -> Path:
    log("Writing DOCX summary...")
    ensure_output_dir()

    doc = Document()
    today = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Daily Summary – {today}", level=1)
    doc.add_paragraph(summary)
    doc.save(DOCX_PATH)

    log(f"DOCX summary written to {DOCX_PATH}")
    return DOCX_PATH


def append_excel_log(summary: str) -> Path:
    log("Appending to Excel log...")
    ensure_output_dir()

    if XLSX_PATH.exists():
        workbook = load_workbook(XLSX_PATH)
        worksheet = workbook.active
    else:
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Daily Log"
        worksheet.append(["Timestamp", "Summary"])

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    worksheet.append([timestamp, summary])
    workbook.save(XLSX_PATH)

    log(f"Excel log updated at {XLSX_PATH}")
    return XLSX_PATH


def write_json_summary(summary_text: str, path: Path) -> Path:
    """Write a minimal JSON payload for the daily summary."""
    today = datetime.now().strftime("%Y-%m-%d")

    payload = {
        "date": today,
        "summary_text": summary_text,
        "raw_text": summary_text,
        "created_at": datetime.now().isoformat(timespec="seconds"),
    }

    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as file_handle:
        json.dump(payload, file_handle, ensure_ascii=False, indent=2)

    log(f"JSON summary written to {path}")
    return path


def get_github_session():
    """Return a configured (session, repo) tuple, or (None, None) if missing."""
    token = os.getenv("GITHUB_TOKEN")
    repo = os.getenv("GITHUB_REPO")

    if not token or not repo:
        log("GitHub config not found (GITHUB_TOKEN / GITHUB_REPO). Skipping GitHub step.")
        return None, None

    session = requests.Session()
    session.headers.update(
        {
            "Authorization": f"Bearer {token}",
            "Accept": "application/vnd.github+json",
            "X-GitHub-Api-Version": "2022-11-28",
            "User-Agent": "daily-automation-script",
        }
    )
    return session, repo


def create_or_update_github_task(summary: str):
    """
    Create or update a GitHub issue for today's daily automation run.

    - Looks for an open issue with today's title.
    - If found, adds a comment with the latest summary.
    - If not found, creates a new issue with labels.
    """
    session, repo = get_github_session()
    if not session or not repo:
        return None

    today = datetime.now().strftime("%Y-%m-%d")
    title = f"Daily Automation – {today}"
    issues_url = f"https://api.github.com/repos/{repo}/issues"

    body_block = f"""# Daily Automation – {today}

{summary}

---

_Repo: {repo} • Generated by daily_v2.py_
"""

    try:
        response = session.get(issues_url, params={"state": "open", "per_page": 50})
        response.raise_for_status()
        issues = response.json()
        existing = next((issue for issue in issues if issue.get("title") == title), None)

        if existing:
            comments_url = existing.get("comments_url")
            comment_response = session.post(comments_url, json={"body": body_block})
            comment_response.raise_for_status()
            url = existing.get("html_url")
            log(f"Updated existing GitHub issue: {url}")
            return url

        payload = {
            "title": title,
            "body": body_block,
            "labels": ["automation", "daily-run"],
        }
        issue_response = session.post(issues_url, json=payload)
        issue_response.raise_for_status()
        issue = issue_response.json()
        url = issue.get("html_url")
        log(f"Created new GitHub issue: {url}")
        return url
    except Exception as exc:  # noqa: BLE001
        log(f"GitHub issue creation/update failed: {exc}")
        return None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Daily automation runner (v2)")
    parser.add_argument(
        "--dry_run",
        action="store_true",
        help="Run without writing files or touching GitHub; log intended actions.",
    )
    return parser.parse_args()


def main(run: bool = True) -> int:
    log(f"=== Daily v2 automation run starting (run={run}) ===")

    client = check_openai_client()
    notes_text = load_recent_notes()
    summary_text = summarize_notes(client, notes_text)

    if not run:
        log("DRY RUN: skipping file writes and GitHub interactions.")
        log("Summary preview:")
        log(summary_text)
        log("=== Daily v2 automation run finished (dry run) ===")
        return 0

    txt_path = write_txt_summary(summary_text)
    docx_path = write_docx_summary(summary_text)
    xlsx_path = append_excel_log(summary_text)

    json_path = write_json_summary(summary_text, JSON_PATH)
    write_json_summary(summary_text, REPO_JSON_PATH)

    issue_url = create_or_update_github_task(summary_text)

    log("Run complete.")
    log(
        "Outputs:\n"
        f"  TXT:  {txt_path}\n"
        f"  DOCX: {docx_path}\n"
        f"  XLSX: {xlsx_path}\n"
        f"  JSON: {json_path}\n"
        f"  Repo JSON: {REPO_JSON_PATH}"
    )
    if issue_url:
        log(f"GitHub task: {issue_url}")
    else:
        log("GitHub task: skipped or failed (see logs above).")

    log("=== Daily v2 automation run finished successfully ===")
    return 0


if __name__ == "__main__":
    args = parse_args()
    sys.exit(main(run=not args.dry_run))
